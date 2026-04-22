import streamlit as st
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import csv
import spacy
import os
import anthropic
from dotenv import load_dotenv

# Load API key
load_dotenv()
api_key = os.getenv("ANTHROPIC_API_KEY")

# ── PAGE CONFIG ────────────────────────────────────────────
st.set_page_config(
    page_title="Contract Risk Analyzer",
    page_icon="📋",
    layout="wide"
)

# ── CUSTOM STYLING ─────────────────────────────────────────
st.markdown("""
<style>
    .main { background-color: #f8f9fa; }
    .block-container { padding-top: 1.5rem; }
    h1 { color: #1F4E79; font-family: Arial; }
    h2, h3 { color: #2E75B6; font-family: Arial; }
    .section-divider {
        border-top: 2px solid #2E75B6;
        margin: 20px 0px;
    }
    .footer {
        text-align: center;
        color: #888888;
        font-size: 12px;
        margin-top: 40px;
        font-family: Arial;
    }
</style>
""", unsafe_allow_html=True)

# ── ACCESS CODE FOR AI SUGGESTIONS ────────────────────────
def check_access_code():
    st.sidebar.markdown("---")
    st.sidebar.markdown("### AI Suggestions")
    code = st.sidebar.text_input(
        "Enter access code for AI suggestions:",
        type="password",
        help="Contact Louis Filiano for access code"
    )
    if code == "Birthday-41":
        st.sidebar.success("AI suggestions unlocked!")
        return True
    elif code != "":
        st.sidebar.error("Incorrect access code")
        return False
    else:
        st.sidebar.info(
            "Enter access code to enable "
            "AI-powered clause suggestions"
        )
        return False

# ── LOAD SPACY ─────────────────────────────────────────────
@st.cache_resource
def load_nlp():
    return spacy.load("en_core_web_sm")

nlp = load_nlp()

# ── CONTRACT TYPE CONFIGURATION ────────────────────────────
CONTRACT_TYPES = {
    "Sterilization Services Agreement": {
        "library": "clause_library_sterilization.csv",
        "context": "sterilization services for regulated manufacturing",
        "critical_missing": [
            "Parties & Legal Entities", "Term & Renewal", "Scope of Supply",
            "Product Specifications", "Pricing", "Delivery Terms",
            "Payment Terms", "Invoicing",
            "Quality Standards", "Warranties", "Inspection & Acceptance",
            "Compliance with Laws", "Regulatory Requirements", "Insurance",
            "Indemnification", "Liability Cap", "Product Recall",
            "Confidentiality", "Intellectual Property",
            "Termination for Convenience", "Termination for Cause", "Effect of Termination",
            "Force Majeure", "Business Continuity",
            "Governing Law", "Dispute Resolution", "Assignment",
            "Capacity Commitment", "Sterilization Certificate", "Validation Protocol",
            "Dose Mapping", "ETO Residual Limits", "Facility Change Control", "Parent Guarantee"
        ],
        "clause_groups": {
            "Basic Contract Setup": ["Parties & Legal Entities", "Term & Renewal", "Scope of Supply", "Definitions"],
            "Products & Delivery": ["Product Specifications", "Pricing", "Delivery Terms", "Title & Risk of Loss", "Forecasting & Ordering", "Packaging & Labeling"],
            "Payment & Financial": ["Payment Terms", "Invoicing", "Taxes & Duties", "Currency"],
            "Quality & Performance": ["Quality Standards", "Warranties", "Inspection & Acceptance", "Remedies & Cure"],
            "Compliance & Responsibility": ["Compliance with Laws", "Regulatory Requirements", "Ethics & Code of Conduct", "Insurance"],
            "Liability & Risk": ["Indemnification", "Liability Cap", "Product Recall"],
            "Confidentiality & Data": ["Confidentiality", "Data Protection & Privacy", "Intellectual Property"],
            "Termination & Exit": ["Termination for Convenience", "Termination for Cause", "Effect of Termination"],
            "Force Majeure & Continuity": ["Force Majeure", "Business Continuity"],
            "General Boilerplate": ["Governing Law", "Dispute Resolution", "Assignment", "Subcontracting", "Amendments & Modifications", "Notices", "Entire Agreement"],
            "Sterilization Specific": ["Capacity Commitment", "Sterilization Certificate", "Validation Protocol", "Dose Mapping", "ETO Residual Limits", "Facility Change Control", "Parent Guarantee", "LLC Structure Risk", "Covered Entity Scope"]
        }
    },
    "MRO / Indirect Spend Agreement": {
        "library": "clause_library_mro.csv",
        "context": "MRO and indirect spend procurement",
        "critical_missing": [
            "Parties & Legal Entities", "Term & Renewal", "Scope of Supply",
            "Product Specifications", "Pricing", "Delivery Terms",
            "Payment Terms", "Invoicing", "Quality Standards", "Warranties",
            "Compliance with Laws", "Insurance",
            "Indemnification", "Liability Cap",
            "Confidentiality", "Termination for Convenience", "Termination for Cause",
            "Force Majeure", "Business Continuity", "Governing Law", "Dispute Resolution"
        ],
        "clause_groups": {
            "Basic Contract Setup": ["Parties & Legal Entities", "Term & Renewal", "Scope of Supply", "Definitions"],
            "Products & Delivery": ["Product Specifications", "Pricing", "Delivery Terms", "Title & Risk of Loss", "Forecasting & Ordering"],
            "Payment & Financial": ["Payment Terms", "Invoicing", "Taxes & Duties"],
            "Quality & Performance": ["Quality Standards", "Warranties", "Inspection & Acceptance"],
            "Compliance & Responsibility": ["Compliance with Laws", "Insurance"],
            "Liability & Risk": ["Indemnification", "Liability Cap"],
            "Confidentiality & Data": ["Confidentiality", "Intellectual Property"],
            "Termination & Exit": ["Termination for Convenience", "Termination for Cause", "Effect of Termination"],
            "Force Majeure & Continuity": ["Force Majeure", "Business Continuity"],
            "General Boilerplate": ["Governing Law", "Dispute Resolution", "Assignment", "Subcontracting", "Entire Agreement"]
        }
    },
    "Supplier MSA (Master Service Agreement)": {
        "library": "clause_library_msa.csv",
        "context": "master service agreements for supplier partnerships",
        "critical_missing": [
            "Parties & Legal Entities", "Term & Renewal", "Scope of Supply",
            "Payment Terms", "Invoicing",
            "Quality Standards", "Compliance with Laws", "Insurance",
            "Indemnification", "Liability Cap",
            "Confidentiality", "Data Protection & Privacy", "Intellectual Property",
            "Termination for Convenience", "Termination for Cause", "Effect of Termination",
            "Force Majeure", "Business Continuity", "Governing Law", "Dispute Resolution"
        ],
        "clause_groups": {
            "Basic Contract Setup": ["Parties & Legal Entities", "Term & Renewal", "Scope of Supply", "Definitions"],
            "Payment & Financial": ["Payment Terms", "Invoicing"],
            "Compliance & Responsibility": ["Compliance with Laws", "Ethics & Code of Conduct", "Insurance"],
            "Liability & Risk": ["Indemnification", "Liability Cap"],
            "Confidentiality & Data": ["Confidentiality", "Data Protection & Privacy", "Intellectual Property"],
            "Termination & Exit": ["Termination for Convenience", "Termination for Cause", "Effect of Termination"],
            "Force Majeure & Continuity": ["Force Majeure", "Business Continuity"],
            "General Boilerplate": ["Governing Law", "Dispute Resolution", "Assignment", "Amendments & Modifications", "Entire Agreement"]
        }
    },
    "General Procurement Agreement": {
        "library": "clause_library_general.csv",
        "context": "general procurement agreements",
        "critical_missing": [
            "Parties & Legal Entities", "Term & Renewal", "Scope of Supply",
            "Product Specifications", "Pricing", "Delivery Terms",
            "Payment Terms", "Invoicing", "Quality Standards", "Warranties",
            "Compliance with Laws", "Insurance",
            "Indemnification", "Liability Cap",
            "Confidentiality", "Termination for Convenience", "Termination for Cause",
            "Force Majeure", "Governing Law", "Dispute Resolution", "Assignment"
        ],
        "clause_groups": {
            "Basic Contract Setup": ["Parties & Legal Entities", "Term & Renewal", "Scope of Supply", "Definitions"],
            "Products & Delivery": ["Product Specifications", "Pricing", "Delivery Terms", "Title & Risk of Loss"],
            "Payment & Financial": ["Payment Terms", "Invoicing", "Taxes & Duties"],
            "Quality & Performance": ["Quality Standards", "Warranties", "Inspection & Acceptance"],
            "Compliance & Responsibility": ["Compliance with Laws", "Insurance"],
            "Liability & Risk": ["Indemnification", "Liability Cap"],
            "Confidentiality & Data": ["Confidentiality", "Intellectual Property"],
            "Termination & Exit": ["Termination for Convenience", "Termination for Cause", "Effect of Termination"],
            "Force Majeure & Continuity": ["Force Majeure", "Business Continuity"],
            "General Boilerplate": ["Governing Law", "Dispute Resolution", "Assignment", "Subcontracting", "Entire Agreement"]
        }
    }
}

# ── WEIGHTS & SCORING ──────────────────────────────────────
CATEGORY_WEIGHTS = {
    # Existing sterilization-specific
    "Auto-Renewal": 9, "Liability Cap": 10,
    "Indemnification": 9, "Payment Terms": 8,
    "Price Escalation": 9, "Termination Rights": 10,
    "Force Majeure": 7, "Intellectual Property": 8,
    "Exclusivity": 8, "Governing Law": 6,
    "Dose Mapping": 9, "Bioburden Management": 8,
    "Sterilization Certificate": 8, "Capacity Commitment": 10,
    "Business Continuity": 10, "Regulatory Change": 8,
    "Validation Protocol": 10, "ETO Residual Limits": 9,
    "Facility Change Control": 9, "Parent Guarantee": 9,
    "LLC Structure Risk": 9, "Covered Entity Scope": 8,
    "Regulatory Action by Entity": 8, "Insurance by Entity": 7,
    # New standard clauses
    "Parties & Legal Entities": 8,
    "Definitions": 5,
    "Term & Renewal": 8,
    "Scope of Supply": 9,
    "Product Specifications": 9,
    "Forecasting & Ordering": 7,
    "Pricing": 9,
    "Delivery Terms": 8,
    "Title & Risk of Loss": 7,
    "Packaging & Labeling": 6,
    "Invoicing": 7,
    "Taxes & Duties": 6,
    "Currency": 5,
    "Quality Standards": 9,
    "Inspection & Acceptance": 8,
    "Warranties": 9,
    "Remedies & Cure": 8,
    "Compliance with Laws": 9,
    "Regulatory Requirements": 9,
    "Ethics & Code of Conduct": 6,
    "Insurance": 8,
    "Product Recall": 9,
    "Confidentiality": 9,
    "Data Protection & Privacy": 9,
    "Termination for Convenience": 8,
    "Termination for Cause": 9,
    "Effect of Termination": 7,
    "Dispute Resolution": 7,
    "Assignment": 7,
    "Subcontracting": 7,
    "Amendments & Modifications": 6,
    "Notices": 5,
    "Entire Agreement": 6,
    # Legacy
    "Delivery & Lead Time": 9, "Warranty": 8,
    "Inventory & Availability": 8, "Pricing & Catalog": 7,
    "Scope of Services": 8, "Service Levels": 9,
    "Data Protection": 9,
}

RISK_LEVEL_MULTIPLIER = {
    "Risky": 1.0, "Acceptable": 0.3, "Preferred": 0.0}

MISSING_CLAUSE_PENALTY = {
    # Sterilization specific
    "Business Continuity": 15, "Capacity Commitment": 12,
    "Validation Protocol": 12, "Sterilization Certificate": 10,
    "Dose Mapping": 8, "ETO Residual Limits": 8,
    "Facility Change Control": 10, "Parent Guarantee": 10,
    "LLC Structure Risk": 8, "Covered Entity Scope": 8,
    # Standard contract
    "Parties & Legal Entities": 10,
    "Term & Renewal": 8,
    "Scope of Supply": 10,
    "Product Specifications": 10,
    "Pricing": 10,
    "Delivery Terms": 9,
    "Title & Risk of Loss": 7,
    "Forecasting & Ordering": 7,
    "Packaging & Labeling": 5,
    "Payment Terms": 10,
    "Invoicing": 7,
    "Taxes & Duties": 6,
    "Currency": 5,
    "Quality Standards": 10,
    "Warranties": 10,
    "Inspection & Acceptance": 9,
    "Remedies & Cure": 8,
    "Compliance with Laws": 10,
    "Regulatory Requirements": 10,
    "Ethics & Code of Conduct": 6,
    "Insurance": 9,
    "Indemnification": 10,
    "Liability Cap": 12,
    "Product Recall": 10,
    "Confidentiality": 10,
    "Data Protection & Privacy": 10,
    "Intellectual Property": 8,
    "Termination for Convenience": 9,
    "Termination for Cause": 10,
    "Effect of Termination": 7,
    "Force Majeure": 8,
    "Governing Law": 7,
    "Dispute Resolution": 8,
    "Assignment": 7,
    "Subcontracting": 7,
    "Amendments & Modifications": 6,
    "Notices": 5,
    "Entire Agreement": 6,
    # Legacy
    "Termination Rights": 12, "Price Escalation": 7,
    "Delivery & Lead Time": 10, "Warranty": 8,
    "Inventory & Availability": 8, "Pricing & Catalog": 7,
    "Scope of Services": 9, "Service Levels": 10,
    "Data Protection": 10,
}

MIN_SENTENCE_LENGTH = {
    "Force Majeure": 60, "Governing Law": 50,
    "Payment Terms": 40, "Auto-Renewal": 50, "default": 25,
}

EXCLUSION_PHRASES = {
    "Force Majeure": ["invoice", "payment", "net 30",
                      "net 60", "net 90"],
    "Payment Terms": ["terminate", "renew",
                      "liability", "arbitration"],
    "Auto-Renewal":  ["invoice", "payment", "net 30"],
    "Governing Law": ["invoice", "payment"],
}

# ── CORE FUNCTIONS ─────────────────────────────────────────
@st.cache_data
def load_clause_library(filepath):
    """Load clause library — supports both old (4-col) and new (7-col) format"""
    library = {}
    with open(filepath, "r", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        for row in reader:
            category = row["Risk_Category"]
            if category not in library:
                library[category] = []
            library[category].append({
                "risk_level":     row["Risk_Level"],
                "trigger_phrase": row["Trigger_Phrase"].lower(),
                "example":        row["Example_Clause"]
            })
    return library

def split_sentences(text):
    doc = nlp(text)
    return [sent.text.strip() for sent in doc.sents
            if len(sent.text.strip()) > 20]

def detect_clauses(sentences, library):
    findings = []
    for sentence in sentences:
        sentence_lower = sentence.lower()
        for category, clauses in library.items():
            min_len = MIN_SENTENCE_LENGTH.get(
                category, MIN_SENTENCE_LENGTH["default"])
            if len(sentence) < min_len:
                continue
            exclusions = EXCLUSION_PHRASES.get(category, [])
            if any(excl in sentence_lower for excl in exclusions):
                continue
            for clause in clauses:
                trigger = clause["trigger_phrase"]
                if trigger in sentence_lower:
                    findings.append({
                        "sentence":      sentence,
                        "category":      category,
                        "risk_level":    clause["risk_level"],
                        "trigger_found": trigger,
                    })
                    break
    return findings

def check_missing_clauses(findings, contract_config):
    found_categories = set(f["category"] for f in findings)
    return [cat for cat in contract_config["critical_missing"]
            if cat not in found_categories]

def calculate_risk_score(findings, missing_clauses):
    total_score = 0
    breakdown = {}
    for f in findings:
        weight = CATEGORY_WEIGHTS.get(f["category"], 5)
        mult = RISK_LEVEL_MULTIPLIER.get(f["risk_level"], 0)
        pts = weight * mult
        total_score += pts
        if pts > 0:
            breakdown[f["category"]] = \
                breakdown.get(f["category"], 0) + pts
    for cat in missing_clauses:
        penalty = MISSING_CLAUSE_PENALTY.get(cat, 5)
        total_score += penalty
        breakdown[f"MISSING: {cat}"] = penalty
    max_possible = sum(CATEGORY_WEIGHTS.values()) + \
                   sum(MISSING_CLAUSE_PENALTY.values())
    score = min(100, round((total_score / max_possible) * 100))
    if score <= 25:
        label = "LOW RISK"
        color = "#375623"
        action = "Proceed with standard review and signature"
    elif score <= 50:
        label = "MODERATE RISK"
        color = "#FF8C00"
        action = "Address flagged clauses before signing"
    elif score <= 75:
        label = "HIGH RISK"
        color = "#C00000"
        action = "Significant negotiation required before signing"
    else:
        label = "CRITICAL RISK"
        color = "#7B0000"
        action = "Do not sign — escalate to legal immediately"
    return score, label, color, action, breakdown

# ── AI SUGGESTION FUNCTIONS ────────────────────────────────
def get_ai_suggestion(finding, context):
    try:
        client = anthropic.Anthropic(api_key=api_key)
        prompt = f"""You are an expert procurement attorney
specializing in {context} agreements for regulated manufacturing.

A clause has been flagged as RISKY: {finding['category']}
Clause: "{finding['sentence']}"
Trigger: "{finding['trigger_found']}"

Provide:
RISK EXPLANATION: [2 sentences why this is risky]
SUGGESTED LANGUAGE: [improved clause protecting the buyer]"""
        message = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=400,
            messages=[{"role": "user", "content": prompt}]
        )
        return message.content[0].text
    except Exception as e:
        return f"AI suggestion unavailable: {e}"

def get_missing_suggestion(category, context):
    try:
        client = anthropic.Anthropic(api_key=api_key)
        prompt = f"""You are an expert procurement attorney
specializing in {context} agreements for regulated manufacturing.

Critical clause MISSING from contract: {category}

Provide:
WHY IT MATTERS: [2 sentences why this clause is critical]
RECOMMENDED LANGUAGE: [actual clause language to add]"""
        message = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=400,
            messages=[{"role": "user", "content": prompt}]
        )
        return message.content[0].text
    except Exception as e:
        return f"AI suggestion unavailable: {e}"

# ── HEADER ─────────────────────────────────────────────────
st.markdown(
    "<h1 style='text-align:center;'>📋 Contract Risk Analyzer</h1>",
    unsafe_allow_html=True)
st.markdown(
    "<p style='text-align:center; color:#555; font-family:Arial;'>"
    "AI-Powered Contract Clause Detection & Risk Scoring  |  "
    "Louis Filiano</p>", unsafe_allow_html=True)
st.markdown("<div class='section-divider'></div>",
            unsafe_allow_html=True)

# ── SIDEBAR CONTROLS ───────────────────────────────────────
st.sidebar.markdown("## Contract Settings")

contract_type = st.sidebar.selectbox(
    "Contract Type",
    list(CONTRACT_TYPES.keys())
)

uploaded_file = st.sidebar.file_uploader(
    "Upload Contract",
    type=["txt", "docx", "pdf"],
    help="Upload your contract as a Word document (.docx), PDF, or plain text (.txt)"
)

ai_enabled = check_access_code()

st.sidebar.markdown("---")
st.sidebar.markdown("### About This Tool")
st.sidebar.markdown("""
Detects risky clauses across:
- Auto-Renewal & Evergreen terms
- Liability Caps & Indemnification
- Payment Terms & Price Escalation
- Termination Rights & Force Majeure
- Contract-specific categories
""")

# ── MAIN CONTENT ───────────────────────────────────────────
if uploaded_file is not None:

    # Extract text based on file type
    file_ext = uploaded_file.name.split(".")[-1].lower()

    if file_ext == "txt":
        contract_text = uploaded_file.read().decode("utf-8")

    elif file_ext == "docx":
        try:
            from docx import Document
            import io
            doc = Document(io.BytesIO(uploaded_file.read()))
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            contract_text = " ".join(paragraphs)
        except ImportError:
            st.error("python-docx is required for Word documents. Run: pip install python-docx")
            st.stop()
        except Exception as e:
            st.error(f"Could not read Word document: {e}")
            st.stop()

    elif file_ext == "pdf":
        try:
            import pdfplumber
            import io
            pages = []
            with pdfplumber.open(io.BytesIO(uploaded_file.read())) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
            contract_text = " ".join(pages)
        except ImportError:
            st.error("pdfplumber is required for PDF files. Run: pip install pdfplumber")
            st.stop()
        except Exception as e:
            st.error(f"Could not read PDF: {e}")
            st.stop()

    else:
        st.error(f"Unsupported file type: {file_ext}")
        st.stop()

    if not contract_text.strip():
        st.error("Could not extract text from the uploaded file. Please check the file is not empty or password protected.")
        st.stop()

    st.sidebar.success(f"✅ Loaded: {uploaded_file.name} ({len(contract_text):,} characters)")
    contract_config = CONTRACT_TYPES[contract_type]

    with st.spinner("Loading clause library..."):
        library = load_clause_library(contract_config["library"])

    with st.spinner("Analyzing contract..."):
        sentences = split_sentences(contract_text)
        findings = detect_clauses(sentences, library)
        missing = check_missing_clauses(findings, contract_config)
        score, label, color, action, breakdown = \
            calculate_risk_score(findings, missing)

    risky =      [f for f in findings if f["risk_level"] == "Risky"]
    acceptable = [f for f in findings if f["risk_level"] == "Acceptable"]

    # ── SECTION 1: EXECUTIVE SUMMARY ──────────────────────
    st.markdown("## Executive Summary")

    commercial_risky = [f for f in risky if f.get("term_type") == "Commercial"]
    legal_risky = [f for f in risky if f.get("term_type") == "Legal"]

    c1, c2, c3, c4, c5, c6, c7 = st.columns(7)
    c1.metric("Contract Type", contract_type.split("(")[0].strip()[:20])
    c2.metric("Sentences Analyzed", len(sentences))
    c3.metric("Risky Clauses", len(risky))
    c4.metric("🟦 Commercial", len(commercial_risky))
    c5.metric("⚖️ Legal", len(legal_risky))
    c6.metric("Missing Critical", len(missing))
    c7.metric("Risk Score", f"{score}/100")

    st.markdown("### Overall Risk Score")
    col_score, col_action = st.columns([1, 2])

    with col_score:
        fig, ax = plt.subplots(figsize=(4, 3))
        ax.barh(["Risk Score"], [score],
                color=color, height=0.4)
        ax.barh(["Risk Score"], [100 - score],
                left=[score], color="#EEEEEE", height=0.4)
        ax.set_xlim(0, 100)
        ax.set_xlabel("Score (0-100)")
        ax.set_title(f"{score}/100 — {label}",
                     color=color, fontweight="bold")
        ax.axvline(x=25, color="green",
                   linestyle="--", alpha=0.5, linewidth=1)
        ax.axvline(x=50, color="orange",
                   linestyle="--", alpha=0.5, linewidth=1)
        ax.axvline(x=75, color="red",
                   linestyle="--", alpha=0.5, linewidth=1)
        plt.tight_layout()
        st.pyplot(fig)

    with col_action:
        st.markdown(f"**Risk Level:** "
                    f"<span style='color:{color}'>{label}</span>",
                    unsafe_allow_html=True)
        st.markdown(f"**Recommended Action:** {action}")
        st.markdown("**Score Guide:**")
        st.markdown("🟢 0-25: Low Risk — proceed with review")
        st.markdown("🟡 26-50: Moderate — address before signing")
        st.markdown("🔴 51-75: High — negotiate before signing")
        st.markdown("🚨 76-100: Critical — do not sign")

    st.markdown("<div class='section-divider'></div>",
                unsafe_allow_html=True)

    # ── SECTION 2: SCORE BREAKDOWN ─────────────────────────
    st.markdown("## Risk Score Breakdown")
    sorted_bd = sorted(breakdown.items(),
                       key=lambda x: x[1], reverse=True)

    fig2, ax2 = plt.subplots(figsize=(12, 5))
    colors_bd = ["#C00000" if "MISSING" in cat else "#2E75B6"
                 for cat, _ in sorted_bd]
    ax2.barh([x[0] for x in sorted_bd],
             [x[1] for x in sorted_bd],
             color=colors_bd)
    ax2.set_xlabel("Risk Points")
    ax2.set_title(
        "Risk Contribution by Category  "
        "(Red = Missing Clause  Blue = Found Clause)")
    plt.tight_layout()
    st.pyplot(fig2)

    st.markdown("<div class='section-divider'></div>",
                unsafe_allow_html=True)

    # ── SECTION 3: RISKY CLAUSES ───────────────────────────
    st.markdown("## Risky Clauses Found")

    if len(risky) > 0:
        v1, v2 = st.columns(2)
        v1.metric("Total Risky Clauses", len(risky))
        v2.metric("Categories Affected",
                  len(set(f["category"] for f in risky)))

        for i, f in enumerate(risky, 1):
            with st.expander(
                f"[{i}] {f['category']} — "
                f"Trigger: '{f['trigger_found']}'"):
                st.markdown("**Original Clause:**")
                st.warning(f['sentence'])
                # Get Term Type from library if available
                term_type = f.get("term_type", "")
                term_badge = ""
                if term_type == "Commercial":
                    term_badge = "🟦 **Commercial Term**"
                elif term_type == "Legal":
                    term_badge = "⚖️ **Legal Term**"
                clause_group = f.get("clause_group", "")
                st.markdown(
                    f"**Risk Category:** {f['category']}  |  "
                    f"**Weight:** "
                    f"{CATEGORY_WEIGHTS.get(f['category'], 5)}/10  |  "
                    f"{term_badge}")
                if clause_group:
                    st.markdown(f"**Section:** {clause_group}")
                if ai_enabled and api_key:
                    with st.spinner("Getting AI suggestion..."):
                        suggestion = get_ai_suggestion(
                            f, contract_config["context"])
                        st.success(suggestion)
                else:
                    st.info(
                        "Enter access code in sidebar "
                        "to unlock AI-powered suggested language.")
    else:
        st.success("No risky clauses detected!")

    st.markdown("<div class='section-divider'></div>",
                unsafe_allow_html=True)

    # ── SECTION 4: MISSING CLAUSES ─────────────────────────
    st.markdown("## Missing Critical Clauses")

    if len(missing) > 0:
        st.error(
            f"{len(missing)} critical clauses are missing "
            f"from this contract.")

        # Group missing clauses by section
        clause_groups = contract_config.get("clause_groups", {})
        if clause_groups:
            st.markdown("### Missing Clauses by Section")
            for group_name, group_clauses in clause_groups.items():
                missing_in_group = [c for c in missing if c in group_clauses]
                present_in_group = [c for c in group_clauses if c not in missing]
                if not group_clauses:
                    continue
                coverage_pct = int((len(present_in_group) / len(group_clauses)) * 100)
                color = "#16A34A" if coverage_pct == 100 else "#DC2626" if coverage_pct < 50 else "#D97706"
                status_icon = "✅" if coverage_pct == 100 else "⚠️" if coverage_pct >= 50 else "❌"
                with st.expander(f"{status_icon} {group_name} — {coverage_pct}% covered ({len(missing_in_group)} missing)"):
                    if present_in_group:
                        st.markdown("**Present:**")
                        for c in present_in_group:
                            st.markdown(f"  ✅ {c}")
                    if missing_in_group:
                        st.markdown("**Missing:**")
                        for c in missing_in_group:
                            penalty = MISSING_CLAUSE_PENALTY.get(c, 5)
                            st.markdown(f"  ❌ **{c}** (+{penalty} risk points)")
            st.markdown("---")
            st.markdown("### Missing Clause Details")

        for cat in missing:
            penalty = MISSING_CLAUSE_PENALTY.get(cat, 5)
            with st.expander(
                f"MISSING: {cat} "
                f"(+{penalty} risk points)"):
                st.markdown(
                    f"**{cat}** was not detected in this contract.")
                st.markdown(
                    f"This clause is critical for "
                    f"**{contract_type}** agreements.")
                if ai_enabled and api_key:
                    with st.spinner(
                        "Getting recommended language..."):
                        suggestion = get_missing_suggestion(
                            cat, contract_config["context"])
                        st.success(suggestion)
                else:
                    st.info(
                        "Enter access code in sidebar "
                        "to unlock AI-powered recommended language.")
    else:
        st.success("All critical clauses detected!")

    st.markdown("<div class='section-divider'></div>",
                unsafe_allow_html=True)

    # ── SECTION 5: ACCEPTABLE CLAUSES ─────────────────────
    if len(acceptable) > 0:
        st.markdown("## Acceptable Clauses — Monitor")
        for i, f in enumerate(acceptable, 1):
            with st.expander(f"[{i}] {f['category']}"):
                st.markdown("**Clause:**")
                st.success(f['sentence'])
                st.markdown(
                    "This clause is acceptable but could "
                    "be improved in future negotiations.")

        st.markdown("<div class='section-divider'></div>",
                    unsafe_allow_html=True)

    # ── SECTION 6: DOWNLOAD ────────────────────────────────
    st.markdown("## Download Results")

    output_rows = []
    for f in risky:
        output_rows.append({
            "Type": "RISKY CLAUSE",
            "Category": f["category"],
            "Risk Level": f["risk_level"],
            "Trigger Found": f["trigger_found"],
            "Contract Sentence": f["sentence"],
            "AI Suggestion": "Enter access code for AI suggestions"
        })
    for cat in missing:
        output_rows.append({
            "Type": "MISSING CLAUSE",
            "Category": cat,
            "Risk Level": "Critical",
            "Trigger Found": "N/A",
            "Contract Sentence": "Clause not found in contract",
            "AI Suggestion": "Enter access code for AI suggestions"
        })
    for f in acceptable:
        output_rows.append({
            "Type": "ACCEPTABLE CLAUSE",
            "Category": f["category"],
            "Risk Level": f["risk_level"],
            "Trigger Found": f["trigger_found"],
            "Contract Sentence": f["sentence"],
            "AI Suggestion": "N/A"
        })

    if output_rows:
        output_df = pd.DataFrame(output_rows)
        csv_out = output_df.to_csv(index=False,
                                   encoding="utf-8-sig")
        st.download_button(
            label="Download Contract Analysis as CSV",
            data=csv_out,
            file_name="contract_risk_analysis.csv",
            mime="text/csv"
        )

    st.markdown(
        "<div class='footer'>Contract Risk Analyzer  |  "
        "Louis Filiano  |  "
        "Powered by Python, spaCy & Claude AI</div>",
        unsafe_allow_html=True)

else:
    st.info(
        "Upload a contract text file using the sidebar "
        "to begin analysis.")
    st.markdown("### This tool detects:")
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("""
        **Standard Procurement Risks:**
        - Auto-renewal and evergreen clauses
        - Inadequate liability caps
        - One-sided indemnification
        - Unfavorable payment terms
        - Uncapped price escalation
        - Weak termination rights
        - Broad force majeure clauses
        - IP ownership issues
        - Exclusivity traps
        - Unfavorable governing law
        """)
    with col2:
        st.markdown("""
        **Sterilization-Specific Risks:**
        - Missing Business Continuity clause
        - No Capacity Commitment
        - Missing Validation Protocol approval rights
        - No Sterilization Certificate requirements
        - Missing Dose Mapping obligations
        - No ETO Residual Limit guarantees
        - Bioburden management gaps
        - Regulatory Change cost allocation

        **MRO & MSA Risks also detected**
        """)